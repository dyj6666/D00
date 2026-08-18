# -*- coding: utf-8 -*-
"""make_ota_doc.py —— 《D00 OTA 升级系统顶级说明书》生成器

内容模块：content_part1.py（基础/架构/存储/包格式/安全）
          content_part2.py（通道/BOOT/流程/可靠性）
          content_part3.py（上位机/实测/故障排查/FAQ/附录）
渲染：封面/摘要/目录域/页眉页脚页码/多级标题/表格/代码块/提示框/步骤
输出：docs/D00_OTA_System_Overview.docx
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"D:\GIT-SPACE\D00\docs\D00_OTA_System_Overview.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
ACCENT2 = RGBColor(0x2E, 0x74, 0xB5)
GRAY = RGBColor(0x59, 0x59, 0x59)
WARN_BG = "FFF2CC"

doc = Document()

# ---------- 全局样式 ----------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
normal.paragraph_format.space_after = Pt(4)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in [
    ("Heading 1", 16, ACCENT, 14, 6),
    ("Heading 2", 13, ACCENT2, 10, 4),
    ("Heading 3", 11.5, GRAY, 8, 3),
]:
    st = doc.styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.font.bold = True
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True


def set_ea(run, name="微软雅黑"):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def P(text="", style=None, size=None, bold=False, color=None, align=None):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        if size:
            r.font.size = Pt(size)
        r.font.bold = bold
        if color:
            r.font.color.rgb = color
        set_ea(r)
    if align is not None:
        p.alignment = align
    return p


def add_field(paragraph, instr, hint="更新"):
    run = paragraph.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = instr
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = hint
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    run._r.append(f1); run._r.append(it); run._r.append(f2)
    run._r.append(t); run._r.append(f3)
    return run


def render(sec):
    """渲染一个内容项：(type, ...)"""
    typ = sec[0]
    if typ == "h1":
        P(sec[1], style="Heading 1")
    elif typ == "h2":
        P(sec[1], style="Heading 2")
    elif typ == "h3":
        P(sec[1], style="Heading 3")
    elif typ == "p":
        P(sec[1])
    elif typ == "pb":  # 加粗引导 + 正文
        p = doc.add_paragraph()
        r = p.add_run(sec[1]); r.font.bold = True; set_ea(r)
        r = p.add_run(sec[2]); set_ea(r)
    elif typ == "li":
        p = doc.add_paragraph(style="List Bullet")
        if len(sec) > 2 and sec[2]:
            r = p.add_run(sec[2]); r.font.bold = True; set_ea(r)
        r = p.add_run(sec[1]); set_ea(r)
    elif typ == "code":
        p = doc.add_paragraph()
        r = p.add_run(sec[1])
        r.font.name = "Consolas"
        r.font.size = Pt(8)
        set_ea(r, "宋体")
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F2F2F2")
        p._p.get_or_add_pPr().append(shd)
        p.paragraph_format.space_after = Pt(6)
    elif typ == "note":  # 提示框
        p = doc.add_paragraph()
        r = p.add_run("※ " + sec[1])
        r.font.size = Pt(9.5)
        set_ea(r)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), WARN_BG)
        p._p.get_or_add_pPr().append(shd)
        p.paragraph_format.space_after = Pt(6)
    elif typ == "table":
        header, rows = sec[1], sec[2]
        widths = sec[3] if len(sec) > 3 else None
        t = doc.add_table(rows=1, cols=len(header))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = t.rows[0].cells
        for j, h in enumerate(header):
            hdr[j].text = ""
            r = hdr[j].paragraphs[0].add_run(h)
            r.font.bold = True; r.font.size = Pt(9); set_ea(r)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "D9E2F3")
            hdr[j]._tc.get_or_add_tcPr().append(shd)
        for row in rows:
            cells = t.add_row().cells
            for j, v in enumerate(row):
                cells[j].text = ""
                r = cells[j].paragraphs[0].add_run(str(v))
                r.font.size = Pt(9); set_ea(r)
        if widths:
            for j, w in enumerate(widths):
                for row in t.rows:
                    row.cells[j].width = Cm(w)
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    elif typ == "pagebreak":
        doc.add_page_break()


# ================= 封面 =================
for _ in range(4):
    P()
P("D00 嵌入式工业平台", size=20, bold=True, color=ACCENT,
  align=WD_ALIGN_PARAGRAPH.CENTER)
P("固件 OTA 升级系统 · 顶级说明书", size=30, bold=True, color=ACCENT,
  align=WD_ALIGN_PARAGRAPH.CENTER)
P("从零精通：架构 · 协议 · 安全 · 实战 · 排障", size=14, color=GRAY,
  align=WD_ALIGN_PARAGRAPH.CENTER)
P()
P("STM32F407ZGT6 · FreeRTOS · 多通道传输 · 端到端安全", size=11,
  align=WD_ALIGN_PARAGRAPH.CENTER)
P("文档版本：V2.0（极致详解版）    编制：D00 工程团队", size=10, color=GRAY,
  align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ================= 摘要 =================
P("摘  要", style="Heading 1")
P("本文档是 D00 嵌入式工业平台的固件 OTA（Over-The-Air，空中升级）系统完整技术说明书，"
  "面向从零开始接触嵌入式升级体系的读者，力求“看完即精通”。全文共 14 章 + 3 个附录，"
  "以“背景知识 → 系统架构 → 物质基础（存储）→ 协议细节（包格式）→ 安全模型 → 传输通道 → "
  "状态机 → 运行时序 → 可靠性 → 动手实操 → 实测证据 → 边界与排障 → FAQ → 展望”为主线，"
  "逐层展开一套生产级 OTA 系统的全部设计决策与实现细节。")
P("系统以 STM32F407ZGT6 为核心，FreeRTOS 承载业务，LVGL 驱动人机界面；升级链路采用 "
  "BOOT/RUN 双区布局 + 外部 SPI Flash 双槽（下载暂存 ota_dl 与回滚源 img_lib），"
  "以 AES-256-CTR 加密、SHA-256 + ECDSA P-256 双公钥签名、芯片 UID 绑定、构建号防重放、"
  "版本号防回滚构建端到端安全链；支持 TCP / HTTP / 串口 HOSTLINK / CAN（预留）/ YMODEM 救援"
  "五类通道；以 PENDING 启动确认、自动回滚、BACKUP 自愈、断点续传、断电四阶段恢复保证"
  "“任何一步失败都不变砖”。文中所有协议帧、数据结构、状态机、故障码均给出源码级定义，"
  "并附真实设备上的功能与安全实测记录，可直接作为开发、联调、维护与二次开发的一手依据。")
P()
P("关键词：OTA 升级；安全引导；AES-256-CTR；ECDSA；芯片绑定；防重放；防回滚；双区布局；"
  "外部 Flash 回滚源；断点续传；断电保护；状态机；STM32F407；FreeRTOS", bold=True)
doc.add_page_break()

# ================= 目录 =================
P("目录", style="Heading 1")
p = doc.add_paragraph()
add_field(p, 'TOC \\o "1-3" \\h \\z \\u', "（在 Word 中右键 → 更新域 生成目录）")
doc.add_page_break()

# ================= 渲染内容 =================
from content_part1 import SECTIONS as S1
from content_part2 import SECTIONS as S2
from content_part3 import SECTIONS as S3
ALL = S1 + S2 + S3
for sec in ALL:
    render(sec)

# ---------- 页眉页脚 ----------
sec = doc.sections[0]
sec.top_margin = Cm(2.2)
sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(2.4)
sec.right_margin = Cm(2.4)

hdr = sec.header.paragraphs[0]
hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = hdr.add_run("D00 嵌入式工业平台 · 固件 OTA 升级系统顶级说明书")
r.font.size = Pt(8); r.font.color.rgb = GRAY; set_ea(r)

ftr = sec.footer.paragraphs[0]
ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ftr.add_run("第 ")
add_field(ftr, "PAGE", "1")
r = ftr.add_run(" 页 · 共 ")
add_field(ftr, "NUMPAGES", "1")
r = ftr.add_run(" 页")
for run in ftr.runs:
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY

doc.save(OUT)
print("已生成:", OUT, "  章节项:", len(ALL))
