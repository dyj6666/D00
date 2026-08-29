# -*- coding: utf-8 -*-
"""make_docs.py —— 将 UNIT_TESTING_PLAN.md 生成为 Word(.docx) + PDF

用法：python workflow/make_docs.py
输出：docs/UNIT_TESTING_PLAN.docx / docs/UNIT_TESTING_PLAN.pdf
"""
import os
import re

MD = r"D:\GIT-SPACE\D00\docs\UNIT_TESTING_PLAN.md"
OUT_DIR = r"D:\GIT-SPACE\D00\docs"


def parse_md(path):
    """简单 md 解析：返回块列表 [(type, content)] type: h1/h2/h3/code/list/table/p"""
    blocks = []
    with open(path, encoding="utf-8") as f:
        lines = f.read().splitlines()
    i = 0
    in_code = False
    code_buf = []
    while i < len(lines):
        ln = lines[i]
        if ln.strip().startswith("```"):
            if in_code:
                blocks.append(("code", "\n".join(code_buf)))
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(ln)
            i += 1
            continue
        if ln.startswith("### "):
            blocks.append(("h3", ln[4:].strip()))
        elif ln.startswith("## "):
            blocks.append(("h2", ln[3:].strip()))
        elif ln.startswith("# "):
            blocks.append(("h1", ln[2:].strip()))
        elif ln.strip().startswith("- "):
            blocks.append(("li", ln.strip()[2:]))
        elif ln.strip().startswith("|"):
            # 表格行
            cells = [c.strip() for c in ln.strip().strip("|").split("|")]
            blocks.append(("tr", cells))
        elif ln.strip() == "":
            pass
        else:
            blocks.append(("p", ln.strip()))
        i += 1
    if in_code:
        blocks.append(("code", "\n".join(code_buf)))
    return blocks


def make_docx(blocks, out):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    # 默认字体（中文用宋体/雅黑）
    style = doc.styles["Normal"]
    style.font.name = "Consolas"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(
        "{%s}eastAsia" % "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "微软雅黑")

    table_rows = []
    for typ, content in blocks:
        if typ == "h1":
            doc.add_heading(content, level=0)
        elif typ == "h2":
            doc.add_heading(content, level=1)
        elif typ == "h3":
            doc.add_heading(content, level=2)
        elif typ == "p":
            doc.add_paragraph(content)
        elif typ == "li":
            doc.add_paragraph(content, style="List Bullet")
        elif typ == "code":
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
            # 浅灰底纹
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), "F2F2F2")
            p._p.get_or_add_pPr().append(shd)
        elif typ == "tr":
            table_rows.append(content)

    # 表格（最后汇总成一张）
    if table_rows:
        header = table_rows[0]
        t = doc.add_table(rows=1, cols=len(header))
        t.style = "Light Grid Accent 1"
        for j, h in enumerate(header):
            t.rows[0].cells[j].text = h
        for row in table_rows[1:]:
            if len(row) != len(header):
                continue
            cells = t.add_row().cells
            for j, v in enumerate(row):
                cells[j].text = v

    doc.save(out)
    print("Word 已生成:", out)


def make_pdf(blocks, out):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Preformatted, Table, TableStyle)

    # 注册中文字体
    font_path = r"C:\Windows\Fonts\simhei.ttf"
    pdfmetrics.registerFont(TTFont("SimHei", font_path))

    styles = {
        "h1": ParagraphStyle("h1", fontName="SimHei", fontSize=18,
                             leading=24, spaceAfter=8,
                             textColor=colors.HexColor("#1F4E79")),
        "h2": ParagraphStyle("h2", fontName="SimHei", fontSize=14,
                             leading=20, spaceBefore=10, spaceAfter=6,
                             textColor=colors.HexColor("#2E74B5")),
        "h3": ParagraphStyle("h3", fontName="SimHei", fontSize=12,
                             leading=16, spaceBefore=6, spaceAfter=4),
        "p": ParagraphStyle("p", fontName="SimHei", fontSize=10,
                            leading=15, spaceAfter=4),
        "li": ParagraphStyle("li", fontName="SimHei", fontSize=10,
                             leading=15, leftIndent=12, bulletIndent=4,
                             spaceAfter=2),
        "code": ParagraphStyle("code", fontName="SimHei", fontSize=8,
                               leading=11, backColor=colors.HexColor("#F2F2F2"),
                               borderColor=colors.HexColor("#CCCCCC"),
                               borderWidth=0.5, borderPadding=4,
                               spaceAfter=6),
    }

    doc = SimpleDocTemplate(out, pagesize=A4,
                            leftMargin=18*mm, rightMargin=18*mm,
                            topMargin=15*mm, bottomMargin=15*mm,
                            title="Unity/Ceedling 单元测试接入方案")

    story = []
    table_rows = []
    for typ, content in blocks:
        if typ == "h1":
            story.append(Paragraph(content, styles["h1"]))
        elif typ == "h2":
            story.append(Paragraph(content, styles["h2"]))
        elif typ == "h3":
            story.append(Paragraph(content, styles["h3"]))
        elif typ == "p":
            story.append(Paragraph(content.replace("&", "&amp;")
                                   .replace("<", "&lt;").replace(">", "&gt;"),
                                   styles["p"]))
        elif typ == "li":
            story.append(Paragraph("• " + content.replace("&", "&amp;")
                                   .replace("<", "&lt;").replace(">", "&gt;"),
                                   styles["li"]))
        elif typ == "code":
            story.append(Preformatted(content, styles["code"]))
        elif typ == "tr":
            table_rows.append(content)
        story.append(Spacer(1, 2))

    if table_rows:
        header = table_rows[0]
        data = [h.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                for h in header]
        rows = [data]
        for row in table_rows[1:]:
            if len(row) != len(header):
                continue
            rows.append([c.replace("&", "&amp;").replace("<", "&lt;")
                         .replace(">", "&gt;") for c in row])
        t = Table(rows)
        t.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, -1), "SimHei"),
            ("FONTSIZE", (0, 0), (-1, -1), 8.5),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2E74B5")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#BBBBBB")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(t)

    doc.build(story)
    print("PDF 已生成:", out)


def main():
    blocks = parse_md(MD)
    make_docx(blocks, os.path.join(OUT_DIR, "UNIT_TESTING_PLAN.docx"))
    make_pdf(blocks, os.path.join(OUT_DIR, "UNIT_TESTING_PLAN.pdf"))


if __name__ == "__main__":
    main()
